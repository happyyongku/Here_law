from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Dict
from docx import Document
from datetime import datetime
import os
import subprocess
from openai import OpenAI
from fastapi.staticfiles import StaticFiles
from fastapi import FastAPI
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# FastAPI 앱 생성
app = FastAPI()

# 라우터 생성
sojang_router = APIRouter()

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.abspath(os.path.join(CURRENT_DIR, ".."))

# OpenAI 클라이언트 생성 (API 키 필요 시 추가)
OPENAI_API_KEY = os.environ["API_KEY"]
client = OpenAI(api_key=OPENAI_API_KEY)  # 여기에 자신의 OpenAI API 키를 입력하세요.

# 사용자로부터 입력받는 데이터 모델 정의
class UserInfo(BaseModel):
    case_title: str
    plaintiff: str
    plaintiff_address: str
    plaintiff_phone: str
    defendant: str
    defendant_address: str
    defendant_phone: str
    court_name: str
    case_details: str

@sojang_router.post("/generate")
async def generate_legal_document(user_info: UserInfo):
    try:
        # 청구 취지 및 내용 생성
        generated_content = generate_content(user_info.dict())

        # Word 파일로 저장
        doc_filename = f"{user_info.plaintiff}_{user_info.case_title}.docx"
        doc_filepath = save_to_word(user_info.dict(), generated_content, filename=doc_filename)

        # Word 파일을 PDF로 변환
        pdf_filename = doc_filename.replace(".docx", ".pdf")
        pdf_filepath = convert_word_to_pdf(doc_filepath, pdf_filename)

        # 파일의 상대 경로를 반환
        return {
            "pdf_url": f"/files/{pdf_filename}",
            "docx_url": f"/files/{doc_filename}"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# OpenAI를 사용해 문서 생성
def generate_content(user_info: Dict[str, str]) -> Dict[str, str]:
    prompt = f"""

    다음 정보를 바탕으로 한국어 소장에 필요한 청구 취지, 청구 원인, 첨부 서류를 작성해주세요.

    사건명:
    {user_info['case_title']}
    원고:
    {user_info['plaintiff']} (주소: {user_info['plaintiff_address']}, 전화번호: {user_info['plaintiff_phone']})
    피고:
    {user_info['defendant']} (주소: {user_info['defendant_address']}, 전화번호: {user_info['defendant_phone']})
    사건 내용:
    {user_info['case_details']}

    소장은 정식 법률 문서 형식으로 작성되어야 하며, 필요한 모든 법적 요소를 포함해야 합니다.
    작성 형식:
    [청구 취지], [청구 원인], [첨부 서류]

    또한 [청구 취지]의 마지막 세줄은 다음의 형식을 무조건 따라야 합니다:
    1. ...
    
    2. 소송비용은 피고가 부담한다.

    3. 제1항은 가집행할 수 있다. 

    라는 판결을 구합니다.
    
    청구취지는 '라는 판결을 구합니다.'항목 이외에는 존댓말을 쓰면 안됩니다. 나머지 청구 원인에는 무조건 존댓말을 써주세요."""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "당신은 한국의 법률 문서 작성 전문가입니다."},
            {"role": "user", "content": prompt}
        ]
    )

    generated_text = response.choices[0].message.content

    sections = {"청구 취지": "", "청구 원인": "", "첨부 서류": ""}
    current_section = None
    for line in generated_text.split('\n'):
        line = line.strip()
        if line.startswith("[청구 취지]"):
            current_section = "청구 취지"
        elif line.startswith("[청구 원인]"):
            current_section = "청구 원인"
        elif line.startswith("[첨부 서류]"):
            current_section = "첨부 서류"
        elif current_section:
            sections[current_section] += line + '\n'

    return sections

# Word로 저장하는 함수 (폰트 적용)
def save_to_word(user_info: Dict[str, str], generated_content: Dict[str, str], filename="소장.docx"):
    SAVE_DIR = os.path.abspath(os.path.join(BASE_DIR, "docs", filename))
    doc = Document()

    # 제목 추가
    title = doc.add_heading('소      장', level=1)
    title.alignment = 1  # 가운데 정렬
    title_run = title.runs[0]
    title_run.font.name = 'NanumGothic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')
    title_run.font.size = Pt(24)

    # 원고 정보
    p1 = doc.add_paragraph(f"원    고 : {user_info['plaintiff']} (전화번호: {user_info['plaintiff_phone']})")
    run1 = p1.runs[0]
    run1.font.name = 'NanumGothic'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p2 = doc.add_paragraph(f"주    소 : {user_info['plaintiff_address']}")
    run2 = p2.runs[0]
    run2.font.name = 'NanumGothic'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 피고 정보
    p3 = doc.add_paragraph(f"피    고 : {user_info['defendant']} (전화번호: {user_info['defendant_phone']})")
    run3 = p3.runs[0]
    run3.font.name = 'NanumGothic'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p4 = doc.add_paragraph(f"주    소 : {user_info['defendant_address']}")
    run4 = p4.runs[0]
    run4.font.name = 'NanumGothic'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 사건명
    case_name = doc.add_paragraph(f"\n사건명: {user_info['case_title']}")
    run_case_name = case_name.runs[0]
    run_case_name.font.name = 'NanumGothic'
    run_case_name._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 청구 취지
    title_claim = doc.add_heading('청 구  취 지', level=2)
    title_claim.alignment = 1
    title_claim_run = title_claim.runs[0]
    title_claim_run.font.name = 'NanumGothic'
    title_claim_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p5 = doc.add_paragraph(generated_content['청구 취지'])
    run5 = p5.runs[0]
    run5.font.name = 'NanumGothic'
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 청구 원인
    title_reason = doc.add_heading('청 구  원 인', level=2)
    title_reason.alignment = 1
    title_reason_run = title_reason.runs[0]
    title_reason_run.font.name = 'NanumGothic'
    title_reason_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p6 = doc.add_paragraph(generated_content['청구 원인'])
    run6 = p6.runs[0]
    run6.font.name = 'NanumGothic'
    run6._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 첨부 서류
    title_attachment = doc.add_heading('첨 부  서 류', level=2)
    title_attachment.alignment = 1
    title_attachment_run = title_attachment.runs[0]
    title_attachment_run.font.name = 'NanumGothic'
    title_attachment_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    for attachment in generated_content['첨부 서류'].split('\n'):
        if attachment.strip() != '':
            p7 = doc.add_paragraph(f"- {attachment.strip()}")
            run7 = p7.runs[0]
            run7.font.name = 'NanumGothic'
            run7._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 날짜 및 서명
    today = datetime.today().strftime('%Y    .   %m    .   %d    .')
    date_paragraph = doc.add_paragraph(today, style='Normal')
    date_paragraph.alignment = 2  # 오른쪽 정렬
    date_run = date_paragraph.runs[0]
    date_run.font.name = 'NanumGothic'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p8 = doc.add_paragraph("원고           (인)", style='Normal')
    p8.alignment = 2
    run8 = p8.runs[0]
    run8.font.name = 'NanumGothic'
    run8._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    p9 = doc.add_paragraph(user_info['plaintiff'], style='Normal')
    p9.alignment = 2
    run9 = p9.runs[0]
    run9.font.name = 'NanumGothic'
    run9._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # 법원 이름
    court_paragraph = doc.add_paragraph(f"\n{user_info['court_name']} 귀중", style='Normal')
    court_paragraph.alignment = 2
    court_run = court_paragraph.runs[0]
    court_run.font.name = 'NanumGothic'
    court_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothic')

    # Word 파일 저장
    doc.save(SAVE_DIR)

    return SAVE_DIR

# Word 파일을 PDF로 변환하는 함수
def convert_word_to_pdf(doc_filepath: str, pdf_filename: str):
    pdf_filepath = os.path.join(BASE_DIR, "docs", pdf_filename)
    
    # 여기서 'libreoffice' 명령어를 사용하여 Word를 PDF로 변환
    command = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_filepath), doc_filepath]
    subprocess.run(command, check=True)

    return pdf_filepath

# PDF 미리보기를 제공하는 라우트
@sojang_router.get("/preview/{pdf_filename}")
async def preview_pdf(pdf_filename: str):
    pdf_filepath = os.path.abspath(os.path.join(BASE_DIR, "docs", pdf_filename))
    if not os.path.exists(pdf_filepath):
        raise HTTPException(status_code=404, detail="PDF 파일을 찾을 수 없습니다.")
    return FileResponse(pdf_filepath, media_type='application/pdf')

# Word 파일 다운로드를 제공하는 라우트
@sojang_router.get("/download/{doc_filename}")
async def download_word_file(doc_filename: str):
    doc_filepath = os.path.abspath(os.path.join(BASE_DIR, "docs", doc_filename))
    if not os.path.exists(doc_filepath):
        raise HTTPException(status_code=404, detail="Word 파일을 찾을 수 없습니다.")
    return FileResponse(doc_filepath, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=doc_filename)

# 정적 파일 제공 설정 (FastAPI 앱에 설정)
app.mount("/files", StaticFiles(directory=os.path.join(BASE_DIR, "docs")), name="files")
